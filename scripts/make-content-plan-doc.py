#!/usr/bin/env python
"""Generate CONTENT-PLAN.docx from the content plan. Throwaway artifact."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x0D, 0x1B, 0x2A)
CYAN = RGBColor(0x0A, 0x7A, 0x8E)
GREY = RGBColor(0x55, 0x5F, 0x6B)

doc = Document()

# Base styles
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x22, 0x2A, 0x33)

for lvl, size in ((1, 20), (2, 15), (3, 12)):
    st = doc.styles[f"Heading {lvl}"]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = NAVY if lvl == 1 else CYAN
    st.font.bold = True


def add_bullet(text, sub=False):
    p = doc.add_paragraph(style="List Bullet 2" if sub else "List Bullet")
    _emit_runs(p, text)
    return p


def _emit_runs(p, text):
    """Render **bold** spans and `code` spans inside a paragraph."""
    import re
    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.color.rgb = CYAN
        else:
            p.add_run(part)


def para(text="", italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    _emit_runs(p, text)
    for r in p.runs:
        if italic:
            r.italic = True
        if color:
            r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def styled_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            _emit_runs(cells[i].paragraphs[0], str(val))
            for rr in cells[i].paragraphs[0].runs:
                rr.font.size = Pt(10)
    return t


# ---- Title block ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("Growth Lift Studio — Content Plan")
tr.font.size = Pt(26); tr.font.bold = True; tr.font.color.rgb = NAVY
para("Data-driven from GSC (3-mo window ending 2026-07-08) + Semrush ranking "
     "best-practices. Baseline: 287 impressions, 12 clicks, 4.2% CTR, avg "
     "position 18.8 (page 2).", italic=True, color=GREY, space_after=14)

# ---- Signals ----
doc.add_heading("The three signals driving this plan", 2)
add_bullet("**Demand is real but stuck on page 2.** The kitchen post pulls 133 "
           "impressions (46% of the whole site) yet 1 click. Impressions exist; "
           "we're not winning them.")
add_bullet("**A \"PPC\" wording cluster is untapped.** `ppc ads for kitchen remodel "
           "company`, `kitchen remodeling ppc ads`, `kitchen remodel ppc`. Every "
           "post says \"Google Ads,\" never \"PPC.\"")
add_bullet("**A broad head term is surfacing with no owner.** `google ads home "
           "improvement`, `google ads for home renovation company` → a pillar/hub "
           "opportunity that also lifts the 4 niche posts via internal links "
           "(topic clusters).")
para("Niches with 0 impressions (roofing, HVAC, siding) = expansion bets, ranked "
     "below the sure things.", italic=True, color=GREY, space_after=10)

# ---- Priority queue ----
doc.add_heading("Priority queue", 2)
styled_table(
    ["#", "Type", "Working title", "Target cluster", "Effort", "Rationale"],
    [
        ["1", "NEW pillar", "Google Ads for Home Improvement Contractors: 2026 Guide",
         "google ads home improvement", "Med", "Owns broad term + hub linking all 4 niche posts"],
        ["2", "OPTIMIZE", "(kitchen post) add PPC cluster + CTR + meta fix",
         "…ppc…kitchen… cluster", "Low", "Highest ROI on site — demand already there"],
        ["3", "META sweep", "Rewrite meta descriptions on all existing posts",
         "CTR across the board", "Low", "Several are >200 chars and truncate in SERP"],
        ["4", "NEW niche", "Google Ads for Roofing Contractors",
         "net-new", "Med", "Best expansion bet: high job value + national volume"],
        ["5", "NEW niche", "Google Ads for HVAC Contractors",
         "net-new", "Med", "Seasonal + high intent, mirrors the window angle"],
        ["6", "NEW comparison", "Facebook Ads vs Google Ads for Contractors",
         "commercial-investigation", "Med", "Comparison format ranks; links both ad-type posts"],
    ],
)
para("")
p = para("Recommended first drop: #1 + #2 + #3 together")
p.runs[0].bold = True
para("Pillar builds the internal-link scaffolding, kitchen fix cashes existing "
     "impressions, meta sweep lifts CTR everywhere at near-zero cost.", space_after=12)

# ---- Meta plan ----
doc.add_heading("Meta plan", 2)
para("Title tag ≤~58 chars, keyword first; description ~150–160 chars, benefit-led "
     "+ CTA.", italic=True, color=GREY)
para("Note: SEO.tsx appends \" | Growth Lift Studio\" to every title — the raw "
     "titles below exclude that.", italic=True, color=GREY, space_after=8)

doc.add_heading("New posts", 3)


def meta_block(name, title_txt, slug, meta_txt, extra=None):
    p = para(f"**{name}**", space_after=2)
    add_bullet(f"Title: `{title_txt}`")
    add_bullet(f"Slug: `{slug}`")
    add_bullet(f"Meta: “{meta_txt}”")
    if extra:
        add_bullet(extra)


meta_block("Pillar (#1)",
           "Google Ads for Home Improvement Contractors: 2026 Guide",
           "google-ads-home-improvement-contractors",
           "How home improvement contractors use Google Ads to book high-ticket "
           "jobs — budgets, campaign structure, and real cost-per-lead numbers by trade.")
meta_block("Roofing (#4)",
           "Google Ads for Roofing Contractors: The 2026 Guide",
           "google-ads-roofing-contractors",
           "How roofing contractors use Google Ads to land high-ticket replacements "
           "— storm-season timing, keywords that convert, and realistic 2026 budgets.")
meta_block("HVAC (#5)",
           "Google Ads for HVAC Contractors: How to Get Leads",
           "google-ads-hvac-contractors",
           "How HVAC contractors use Google Ads to book installs and service calls — "
           "seasonal demand, emergency-intent keywords, and real cost-per-lead numbers.")
meta_block("Comparison (#6)",
           "Facebook Ads vs Google Ads for Contractors (2026)",
           "facebook-ads-vs-google-ads-contractors",
           "Facebook Ads or Google Ads for your contracting business? A straight "
           "comparison of cost, intent, and speed — and when to run each.")

doc.add_heading("Existing posts — meta rewrites (#3)", 3)
meta_block("kitchen (current desc ~230 chars, truncates)",
           "keep existing title",
           "google-ads-kitchen-remodeling-contractors",
           "Kitchen remodeling Google Ads (PPC) done right — keyword strategy, "
           "negative keywords, and landing pages that book $30k+ renovation jobs. "
           "Real numbers.",
           extra="Also work “PPC” into an H2/body naturally to catch the "
                 "untapped cluster.")
add_bullet("**google-ads-bathroom** — audit length; keep if ~150–160, else trim.")
add_bullet("**facebook-ads-bathroom** — audit length; ensure benefit + CTA.")
add_bullet("**google-ads-window** — audit length; keep seasonal hook.")

# ---- Pillar spec ----
doc.add_heading("Blog post spec — Pillar (#1), ready to write", 2)
para("**Intent:** informational/commercial — a contractor deciding whether/how to "
     "run Google Ads before narrowing to their trade. Answer the broad question, "
     "then route to the trade-specific deep-dive.")
para("**H2 outline** (≥4 H2s so the auto mid-post CTA fires; answer-first / BLUF opening):")
for i, h in enumerate([
    "Why Google Ads works for home improvement contractors (hunting-not-fishing hook)",
    "Google Ads vs Facebook Ads for contractors (links the Facebook post)",
    "What it costs by trade — real 2026 numbers (table: kitchen/bathroom/window CPC, "
    "CPL, job value, ROAS; each row links its niche post)",
    "The 5 pillars of a campaign that books jobs (negative keywords, phrase/exact "
    "match, dedicated landing page, conversion tracking, speed-to-lead)",
    "How much should you budget? (worked ROI math)",
    "The Growth Lift Studio advantage ($1,500 → first 5 appts free → 1.5% of closed job)",
    "Where to start (CTA + “read the guide for your trade” internal links)",
]):
    pp = doc.add_paragraph(style="List Number")
    _emit_runs(pp, h)

para("**Internal links (the point of a pillar):** inline absolute "
     "`https://growthliftstudio.in/blog/...` to all 4 existing posts, and add the "
     "new slug to `relatedPostsMap` in both directions.")
para("**Images (IMAGE:: placeholders — generate PNGs into `public/images/` later):**")
add_bullet("`IMAGE::/images/home-improvement-hub-diagram.png::How Google Ads maps "
           "to each home improvement trade`")
add_bullet("`IMAGE::/images/home-improvement-cost-by-trade.png::Google Ads "
           "cost-per-lead and ROAS by trade — 2026`")

# ---- Wiring checklist ----
doc.add_heading("Wiring checklist (per new post — all three or it's invisible)", 2)
for h in [
    "Object in `src/data/blogPosts.ts`",
    "`<url>` in `public/sitemap.xml` (single source of truth for prerender)",
    "Entry in `relatedPostsMap` in `src/pages/BlogPost.tsx` (both directions)",
]:
    pp = doc.add_paragraph(style="List Number")
    _emit_runs(pp, h)
para("**Verify:** `npm run lint` → `npm run build` → check prerendered "
     "`dist/blog/<slug>/index.html` for title/meta/canonical + "
     "`\"@type\":\"BlogPosting\"` + H1; confirm no leaked `**` (renderer only "
     "supports full-line bold).")

# ---- Sequencing ----
doc.add_heading("Sequencing", 2)
add_bullet("**Drop 1 (now):** #1 pillar + #2 kitchen optimize + #3 meta sweep")
add_bullet("**Drop 2:** #4 roofing")
add_bullet("**Drop 3:** #5 HVAC")
add_bullet("**Drop 4:** #6 comparison")
add_bullet("After each drop: re-pull GSC in 2–4 weeks; watch kitchen position/CTR "
           "and pillar impressions.")

out = "d:/Growth Lift Studio/Social Media Brand/New folder/CONTENT-PLAN.docx"
doc.save(out)
print("Saved:", out)
